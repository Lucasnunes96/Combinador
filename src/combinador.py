import sys
from PyQt5.QtWidgets import QApplication, QMainWindow, QListWidget, QListWidgetItem, QPushButton, \
    QAbstractItemView, QLabel, QFileDialog, QWidget, QComboBox, QLineEdit, QMessageBox, QFrame
from PyQt5.QtCore import Qt
from docx import Document
from docxcompose.composer import Composer
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import logging
from datetime import datetime
import src.functions

class ListboxWidget(QListWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setAcceptDrops(True)
        self.setGeometry(400, 100, 600, 175)
        self.setDefaultDropAction(Qt.MoveAction)
        self.setSelectionMode(QAbstractItemView.SingleSelection)
        self.current_item = None
        self.setUpdatesEnabled(True)


    def dragEnterEvent(self, event):
        if event.mimeData().hasUrls():
            event.accept()
        elif event.mimeData().hasFormat('application/x-qabstractitemmodeldatalist'):
            event.accept()
        else:
            event.ignore()

    def dragMoveEvent(self, event):
        if event.mimeData().hasUrls():
            event.setDropAction(Qt.CopyAction)
            urls = event.mimeData().urls()
            for url in urls:
                if not url.toLocalFile().lower().endswith('.docx'):
                    event.ignore()
                    return
            event.accept()
        elif event.mimeData().hasFormat('application/x-qabstractitemmodeldatalist'):
            event.setDropAction(Qt.MoveAction)
            event.accept()
        else:
            event.ignore()

    def dropEvent(self, event):
        if event.mimeData().hasUrls():
            event.setDropAction(Qt.CopyAction)
            event.accept()

            links = []

            for url in event.mimeData().urls():
                links.append(str(url.toLocalFile()))

            self.addItems(links)


class especif_arquiv(QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setGeometry(10, 10, 500, 500)

        self.tipo_sec_label = QLabel("Tipo de Sessão", self)
        self.tipo_sec_label.setGeometry(0, 0, 100, 25)
        self.escolh_sec = QComboBox(self)
        self.escolh_sec.addItems(['Plenária', '1ª Câmara', '2ª Câmara'])
        self.escolh_sec.setGeometry(110, 0, 100, 25)

        def conferir_data(data):
            try:
                date = datetime.strptime(data, '%d/%m/%Y')
                if not 1 <= date.day <= 31 and not 1 <= date.month <= 12:
                    QMessageBox.warning(self, "Atenção", "Por favor utilize o formato de data dd/mm/aaaa.")
                    self.escolh_data.clear()
                else:
                    pass
            except ValueError:
                QMessageBox.warning(self, "Atenção", "Por favor utilize o formato de data dd/mm/aaaa.")
                self.escolh_data.clear()

        self.escolh_data_label = QLabel("Data da Sessão", self)
        self.escolh_data_label.setGeometry(0, 30, 100, 25)
        self.escolh_data = QLineEdit(self)
        self.escolh_data.setPlaceholderText("Ex.: 25/04/2023")
        self.escolh_data.setGeometry(110, 30, 100, 25)
        self.escolh_data.editingFinished.connect(lambda: conferir_data(self.escolh_data.text()))

        self.presentes_label = QLabel("Presentes na Sessão", self)
        self.presentes_label.setGeometry(0, 60, 310, 25)
        bottom_line = QFrame(self.presentes_label)
        bottom_line.setGeometry(0, self.presentes_label.height() - 1, self.presentes_label.width(), 1)
        bottom_line.setFrameShape(QFrame.HLine)
        bottom_line.setStyleSheet("color: #c0c0c0")
        self.presentes_label.setStyleSheet("font-weight: bold")

        names = ["Anselmo Roberto de Almeida Brito", "Fernando Ribeiro Toledo", "Maria Cleide Costa Beserra",
                 "Otávio Lessa de Geraldo Santos", "Renata Pereira Pires Calheiros", "Rodrigo Siqueira Cavalcante",
                 "Rosa Maria Ribeiro de Albuquerque", "Alberto Pires Alves de Abreu", "Sérgio Ricardo Maciel"]

        self.labels = []
        self.cond_combos = []

        y = 90

        for name in names:
            label = QLabel(name, self)
            label.setGeometry(0, y, 200, 25)
            self.labels.append(label)

            combo = QComboBox(self)
            combo.addItems(['Escolha, se presente', 'Votante', 'Presidente', 'Presidente em exercício', 'Presente'])
            combo.setGeometry(175, y, 135, 25)
            self.cond_combos.append(combo)

            y += 30

        self.mpc_label = QLabel("Membro do MPC", self)
        self.mpc_label.setGeometry(0, 360, 80, 25)
        self.mpc = QComboBox(self)
        self.mpc.addItems(['Stella de Barros Lima Mero Cavalcante', 'Ricardo Schneider Rodrigues', 'Pedro Barbosa Neto',
                           'Rafael Rodrigues de Alcântara', 'Enio Andrade Pimenta', 'Gustavo Henrique Albuquerque Santos'])
        self.mpc.setGeometry(85, 360, 225, 25)

        linha_mpc = QFrame(self)
        linha_mpc.setGeometry(0, 390, 310, 1)
        linha_mpc.setStyleSheet("color: #c0c0c0")
        linha_mpc.setFrameShape(QFrame.HLine)

        self.escolh_assinante_label = QLabel("Assinante", self)
        self.escolh_assinante_label.setGeometry(0, 400, 60, 25)
        self.assinante = QComboBox(self)
        self.assinante.addItems(['Lucas Nunes Aureliano Silva', 'Jéssica Luana Silva de Lima',
                                 'André Henrique da Rocha Alencar Rego'])
        self.assinante.setGeometry(85, 400, 225, 25)

        self.cargo_assin_label = QLabel("Cargo do Assinante", self)
        self.cargo_assin_label.setGeometry(0, 430, 100, 25)
        self.cargo_assin = QLineEdit(self)
        self.cargo_assin.setGeometry(110, 430, 100, 25)

        self.matric_ass_lab = QLabel("Matrícula do Assinante", self)
        self.matric_ass_lab.setGeometry(0, 460, 125, 25)
        self.matric_ass = QLineEdit(self)
        self.matric_ass.setGeometry(110, 460, 100, 25)


class AppDemo(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setFixedSize(1050, 525)
        self.setWindowTitle("Gerenciador de Aposentadorias. Closed Beta v1.3")
        self.lstbox_view = ListboxWidget(self)
        self.secao_view = especif_arquiv(self)

        self.btn_delete = QPushButton('Remover arquivo', self)
        self.btn_delete.setGeometry(400, 295, 200, 50)
        self.btn_delete.clicked.connect(self.delete_item)

        self.directory_btn = QPushButton('Escolha o destino', self)
        self.directory_btn.setGeometry(600, 295, 200, 50)
        self.directory_btn.clicked.connect(self.get_directory)

        self.juntar_btn = QPushButton('Gerar arquivo para publicação', self)
        self.juntar_btn.setGeometry(800, 295, 200, 50)
        self.juntar_btn.clicked.connect(self.juntar_arq)

        self.arraste_aqui = QLabel("Arraste os arquivos na seção abaixo (apenas .docx).", self)
        self.arraste_aqui.setGeometry(550, 65, 500, 50)

        self.destino_label = QLabel("Destino: ", self)
        self.destino_label.setGeometry(400, 260, 75, 50)
        self.escolha_destino = QLabel(self)
        self.escolha_destino.setGeometry(445, 260, 600, 50)

        self.up_btn = QPushButton('↑', self)
        self.up_btn.setGeometry(370, 100, 25, 25)
        self.up_btn.clicked.connect(lambda: src.functions.move_item_up(self.lstbox_view))

        self.down_btn = QPushButton('↓', self)
        self.down_btn.setGeometry(370, 130, 25, 25)
        self.down_btn.clicked.connect(lambda: src.functions.move_item_down(self.lstbox_view))


    def juntar_arq(self, filename_master=None):
        logger = logging.getLogger(__name__)
        logger.setLevel(logging.DEBUG)

        formatter = logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s')
        file_handler = logging.FileHandler('../debug.log')
        file_handler.setLevel(logging.DEBUG)
        file_handler.setFormatter(formatter)
        logger.addHandler(file_handler)

        try:
            logger.info('Começando a função juntar_arq')

            if not self.secao_view.escolh_data.text() or not self.secao_view.matric_ass or not self.secao_view.cargo_assin:
                QMessageBox.warning(self, "Atenção", "Por favor, preencha todos os campos antes de continuar.")
                logger.warning('Campos obrigatórios não preenchidos')
            else:
                items = [self.lstbox_view.item(i).text() for i in range(self.lstbox_view.count())]
                number_of_sections = len(items)
                if filename_master:
                    master = Document(filename_master)
                    logger.info('Arquivo mestre carregado com êxito')
                else:
                    master = Document()
                    logger.info('Criado novo arquivo mestre')
                style = master.styles['Normal']
                font = style.font
                font.name = 'Times New Roman'
                font.size = Pt(12)
                section = master.sections[0]
                section.left_margin = Cm(2)
                section.right_margin = Cm(2)
                section.top_margin = Cm(2)
                section.bottom_margin = Cm(2)
                composer = Composer(master)
                if self.secao_view.escolh_sec.currentText() == 'Plenária':
                    sessao = "PLENÁRIA"
                elif self.secao_view.escolh_sec.currentText() == '1ª Câmara':
                    sessao = "DA PRIMEIRA CÂMARA"
                elif self.secao_view.escolh_sec.currentText() == '2ª Câmara':
                    sessao = "DA SEGUNDA CÂMARA"
                logger.debug('Sessão escolhida: %s', sessao)

                dia = src.functions.mes_do_ano(self.escolh_data.text())

                cabecalho = master.add_paragraph()
                cabecalho.add_run('A CONSELHEIRA SUBSTITUTA DO TRIBUNAL DE CONTAS DE ALAGOAS, ').bold = False
                cabecalho.add_run(f"ANA RAQUEL RIBEIRO SAMPAIO CALHEIROS").bold = True
                cabecalho.add_run(f", NA SESSÃO ").bold = False
                cabecalho.add_run(f"{sessao}").bold = True
                cabecalho.add_run(f" DO DIA ").bold = False
                cabecalho.add_run(f"{dia}").bold = True
                cabecalho.add_run(', relatou os seguintes processos:').bold = False
                logger.debug('cabeçalho')

                sexo = {
                    "Rosa Maria Ribeiro de Albuquerque": "Conselheira",
                    "Renata Pereira Pires Calheiros": "Conselheira",
                    "Maria Cleide Costa Beserra": "Conselheira",
                    "Rodrigo Siqueira Cavalcante": "Conselheiro",
                    "Otávio Lessa de Geraldo Santos": "Conselheiro",
                    "Fernando Ribeiro Toledo": "Conselheiro",
                    "Anselmo Roberto de Almeida Brito": "Conselheiro",
                    "Alberto Pires Alves de Abreu": "Conselheiro Substituto",
                    "Sérgio Ricardo Maciel": "Conselheiro Substituto",
                    "Stella de Barros Lima Mero Cavalcante": "Procuradora",
                    "Ricardo Schneider Rodrigues": "Procurador",
                    "Pedro Barbosa Neto": "Procurador",
                    "Rafael Rodrigues de Alcântara": "Procurador",
                    "Enio Andrade Pimenta": "Procurador",
                    "Gustavo Henrique Albuquerque Santos": "Procurador"
                }

                if self.secao_view.anselmo_cond.currentText() == "Presidente":
                    presidente = "Anselmo Roberto de Almeida Brito"
                elif self.secao_view.fernando_cond.currentText() == "Presidente":
                    presidente = "Fernando Ribeiro Toledo"
                elif self.secao_view.maria_cond.currentText() == "Presidente":
                    presidente = "Maria Cleide Costa Beserra"
                elif self.secao_view.otavio_cond.currentText() == "Presidente":
                    presidente = "Otávio Lessa de Geraldo Santos"
                elif self.secao_view.renata_cond.currentText() == "Presidente":
                    presidente = "Renata Pereira Pires Calheiros"
                elif self.secao_view.rodrigo_cond.currentText() == "Presidente":
                    presidente = "Rodrigo Siqueira Cavalcante"
                elif self.secao_view.rosa_cond.currentText() == "Presidente":
                    presidente = "Rosa Maria Ribeiro de Albuquerque"
                elif self.secao_view.alberto_cond.currentText() == "Presidente":
                    presidente = "Alberto Pires Alves de Abreu"
                elif self.secao_view.sergio_cond.currentText() == "Presidente":
                    presidente = "Sérgio Ricardo Maciel"
                else:
                    presidente = False

                if self.secao_view.anselmo_cond.currentText() == "Presidente em exercício":
                    presidente_em_exerc = "Anselmo Roberto de Almeida Brito"
                elif self.secao_view.fernando_cond.currentText() == "Presidente em exercício":
                    presidente_em_exerc = "Fernando Ribeiro Toledo"
                elif self.secao_view.maria_cond.currentText() == "Presidente em exercício":
                    presidente_em_exerc = "Maria Cleide Costa Beserra"
                elif self.secao_view.otavio_cond.currentText() == "Presidente em exercício":
                    presidente_em_exerc = "Otávio Lessa de Geraldo Santos"
                elif self.secao_view.renata_cond.currentText() == "Presidente em exercício":
                    presidente_em_exerc = "Renata Pereira Pires Calheiros"
                elif self.secao_view.rodrigo_cond.currentText() == "Presidente em exercício":
                    presidente_em_exerc = "Rodrigo Siqueira Cavalcante"
                elif self.secao_view.rosa_cond.currentText() == "Presidente em exercício":
                    presidente_em_exerc = "Rosa Maria Ribeiro de Albuquerque"
                elif self.secao_view.alberto_cond.currentText() == "Presidente em exercício":
                    presidente_em_exerc = "Alberto Pires Alves de Abreu"
                elif self.secao_view.sergio_cond.currentText() == "Presidente em exercício":
                    presidente_em_exerc = "Sérgio Ricardo Maciel"
                else:
                    presidente_em_exerc = False

                votantes = []

                if self.secao_view.anselmo_cond.currentText() == "Votante":
                    votantes.append("Anselmo Roberto de Almeida Brito")
                if self.secao_view.fernando_cond.currentText() == "Votante":
                    votantes.append("Fernando Ribeiro Toledo")
                if self.secao_view.maria_cond.currentText() == "Votante":
                    votantes.append("Maria Cleide Costa Beserra")
                if self.secao_view.otavio_cond.currentText() == "Votante":
                    votantes.append("Otávio Lessa de Geraldo Santos")
                if self.secao_view.renata_cond.currentText() == "Votante":
                    votantes.append("Renata Pereira Pires Calheiros")
                if self.secao_view.rodrigo_cond.currentText() == "Votante":
                    votantes.append("Rodrigo Siqueira Cavalcante")
                if self.secao_view.rosa_cond.currentText() == "Votante":
                    votantes.append("Rosa Maria Ribeiro de Albuquerque")
                if self.secao_view.alberto_cond.currentText() == "Votante":
                    votantes.append("Alberto Pires Alves de Abreu")
                if self.secao_view.sergio_cond.currentText() == "Votante":
                    votantes.append("Sérgio Ricardo Maciel")

                presente = []

                if self.secao_view.anselmo_cond.currentText() == "Presente":
                    presente.append("Anselmo Roberto de Almeida Brito")
                if self.secao_view.fernando_cond.currentText() == "Presente":
                    presente.append("Fernando Ribeiro Toledo")
                if self.secao_view.maria_cond.currentText() == "Presente":
                    presente.append("Maria Cleide Costa Beserra")
                if self.secao_view.otavio_cond.currentText() == "Presente":
                    presente.append("Otávio Lessa de Geraldo Santos")
                if self.secao_view.renata_cond.currentText() == "Presente":
                    presente.append("Renata Pereira Pires Calheiros")
                if self.secao_view.rodrigo_cond.currentText() == "Presente":
                    presente.append("Rodrigo Siqueira Cavalcante")
                if self.secao_view.rosa_cond.currentText() == "Presente":
                    presente.append("Rosa Maria Ribeiro de Albuquerque")
                if self.secao_view.alberto_cond.currentText() == "Presente":
                    presente.append("Alberto Pires Alves de Abreu")
                if self.secao_view.sergio_cond.currentText() == "Presente":
                    presente.append("Sérgio Ricardo Maciel")

                if presidente:
                    pres = presidente
                    pres1 = "Presidente"
                else:
                    pres = presidente_em_exerc
                    pres1 = "Presidente em exercício"

                for i in range(number_of_sections):
                    doc_temp = Document(items[i])
                    for paragraph in doc_temp.paragraphs:
                        paragraph.style = doc_temp.styles['Normal']
                        paragraph.paragraph_format.space_before = Pt(0)
                        paragraph.paragraph_format.space_after = Pt(0)
                    composer.append(doc_temp)

                    presentes = master.add_paragraph()
                    presentes.add_run(f"\nConselheira Substituta ").bold = False
                    presentes.add_run(f"Ana Raquel Ribeiro Sampaio Calheiros ").bold = True
                    presentes.add_run(f"- Relatora").bold = False


                    presentes.add_run(f"\n{sexo[pres]} ").bold = False
                    presentes.add_run(f"{pres}").bold = True
                    presentes.add_run(f" - {pres1}").bold = False

                    presentes.add_run(f"\nTomaram parte na votação:").bold = True

                    for votante in votantes:
                        presentes.add_run(f"\n{sexo[votante]} ").bold = False
                        presentes.add_run(f"{votante}").bold = True

                    for presen in presente:
                        presentes.add_run(f"\n{sexo[presen]} ").bold = False
                        presentes.add_run(f"{presen}").bold = True
                        presentes.add_run(f" - Presente").bold = False

                    presentes.add_run(f"\n{sexo[self.secao_view.mpc.currentText()]} ").bold = False
                    presentes.add_run(f"{self.secao_view.mpc.currentText()}").bold = True
                    presentes.add_run(f" - Ministério Público de Contas - Presente").bold = False

                assinante = self.secao_view.assinante.currentText()
                cargo_assin = self.secao_view.cargo_assin.text()
                mat_assin = self.secao_view.matric_ass.text()

                assinatura = master.add_paragraph()
                assinatura.add_run(f"\n{assinante}").bold = True
                assinatura.add_run(f"\n{cargo_assin}\nMatrícula {mat_assin}").bold = False
                assinatura.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                composer.save(f"{self.escolha_destino.text()}\\novo.docx")

                for arquivo in items:
                    doc = Document(arquivo)
                    texto_adicionado = False
                    style = doc.styles['Normal']
                    font = style.font
                    font.name = 'Times New Roman'
                    font.size = Pt(12)
                    for paragraph in doc.paragraphs:
                        if not texto_adicionado:
                            presentes = doc.add_paragraph()
                            presentes.add_run(f"\nConselheira Substituta ").bold = False
                            presentes.add_run(f"Ana Raquel Ribeiro Sampaio Calheiros ").bold = True
                            presentes.add_run(f"- Relatora").bold = False

                            presentes.add_run(f"\n{sexo[pres]} ").bold = False
                            presentes.add_run(f"{pres}").bold = True
                            presentes.add_run(f" - {pres1}").bold = False

                            presentes.add_run(f"\nTomaram parte na votação:").bold = True

                            for votante in votantes:
                                presentes.add_run(f"\n{sexo[votante]} ").bold = False
                                presentes.add_run(f"{votante}").bold = True

                            for presen in presente:
                                presentes.add_run(f"\n{sexo[presen]} ").bold = False
                                presentes.add_run(f"{presen}").bold = True
                                presentes.add_run(f" - Presente").bold = False

                            presentes.add_run(f"\n{sexo[self.secao_view.mpc.currentText()]} ").bold = False
                            presentes.add_run(f"{self.secao_view.mpc.currentText()}").bold = True
                            presentes.add_run(f" - Ministério Público de Contas - Presente").bold = False

                            texto_adicionado = True

                    doc.save(arquivo)

                self.lstbox_view.clear()
                QMessageBox.information(self, "Sucesso", "Arquivo gerado com sucesso!")

        except Exception as e:
            logger.exception("Ocorreu um erro: %s", e)

    def get_directory(self):
        folder = QFileDialog.getExistingDirectory(
            self,
            caption='Selecione o destino'
        )
        self.escolha_destino.setText(folder)

    def delete_item(self):
        for item in self.lstbox_view.selectedItems():
            self.lstbox_view.takeItem(self.lstbox_view.row(item))


if __name__ == "__main__":
    app = QApplication(sys.argv)
    main = AppDemo()
    main.show()
    sys.exit(app.exec())