import sys
from PyQt5.QtWidgets import QApplication, QMainWindow, QListWidget, QListWidgetItem, QPushButton, \
    QAbstractItemView, QLabel, QFileDialog, QWidget, QComboBox, QLineEdit, QMessageBox
from PyQt5.QtCore import Qt, QMimeData, QDataStream, QByteArray, QIODevice, QSize
from PyQt5.QtGui import QDrag, QCursor
from docx import Document
from docxcompose.composer import Composer
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import logging

class ListboxWidget(QListWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setAcceptDrops(True)
        self.setGeometry(400, 100, 600, 175)
        self.setDragEnabled(True)
        self.setDragDropMode(QAbstractItemView.InternalMove)
        self.setDefaultDropAction(Qt.MoveAction)
        self.setSelectionMode(QAbstractItemView.SingleSelection)
        self.current_item = None

    def dragEnterEvent(self, event):
        if event.mimeData().hasUrls():
            event.accept()
        elif event.mimeData().hasFormat('application/x-qabstractitemmodeldatalist'):
            event.accept()
        else:
            event.ignore()

    def dragMoveEvent(self, event):
        if event.mimeData().hasUrls():
            event.setDropAction(Qt.CopyAction)
            urls = event.mimeData().urls()
            for url in urls:
                if not url.toLocalFile().lower().endswith(('.docx', '.doc')):
                    event.ignore()
                    return
            event.accept()
        elif event.mimeData().hasFormat('application/x-qabstractitemmodeldatalist'):
            event.setDropAction(Qt.MoveAction)
            event.accept()
        else:
            event.ignore()

    def dropEvent(self, event):
        if event.mimeData().hasUrls():
            event.setDropAction(Qt.CopyAction)
            event.accept()

            links = []

            for url in event.mimeData().urls():
                if url.isLocalFile():
                    if url.toLocalFile().lower().endswith(('.docx', '.doc')):
                        links.append(str(url.toLocalFile()))
                else:
                    links.append(str(url.toString()))

            self.addItems(links)
        elif event.mimeData().hasFormat('application/x-qabstractitemmodeldatalist'):
            mime_data = event.mimeData()
            bstream = event.mimeData().data('application/x-qabstractitemmodeldatalist')
            data_stream = QDataStream(bstream, QIODevice.ReadOnly)

            row = self.currentRow()
            selected_rows = self.selectedIndexes()

            while not data_stream.atEnd():
                row, column, data = self.decodeData(data_stream)

                if row not in [idx.row() for idx in selected_rows]:
                    self.insertItem(row, QListWidgetItem(data))
                elif row != self.currentRow():
                    # Move dragged item to the new location
                    if self.current_item.listWidget() == self:
                        self.takeItem(self.currentRow())
                    self.insertItem(row, self.current_item)
                    self.setCurrentRow(row)

            event.accept()
        else:
            event.ignore()

    def mousePressEvent(self, event):
        super().mousePressEvent(event)
        if event.button() == Qt.LeftButton:
            self.current_item = self.itemAt(event.pos())

    def mouseMoveEvent(self, event):
        super().mouseMoveEvent(event)

        if not self.current_item:
            return

        mime_data = QMimeData()
        mime_data.setData('application/x-qabstractitemmodeldatalist', self.encodeData())

        drag = QDrag(self)
        drag.setMimeData(mime_data)
        drag.setPixmap(self.current_item.icon().pixmap(self.current_item.icon().actualSize(QSize(50, 50))))
        drag.setHotSpot(event.pos() - self.visualItemRect(self.itemAt(event.pos())).topLeft())

        drop_action = drag.exec_(Qt.MoveAction)

        if drop_action == Qt.MoveAction:
            selected_rows = self.selectedIndexes()
            rows = [idx.row() for idx in selected_rows]
            rows.sort(reverse=True)

            # Obter a posição do cursor durante o evento dropEvent
            cursor_pos = self.mapFromGlobal(QCursor.pos())
            new_row = self.row(self.itemAt(cursor_pos))

            # Ajustar a posição do item antes de inseri-lo na nova posição
            if new_row is None:
                new_row = self.count() - 1
            elif new_row in rows:
                new_row = min(rows)

            for row in rows:
                item = self.takeItem(row)
                self.insertItem(new_row, item)
                new_row += 1

            self.setCurrentRow(rows[-1] + 1)

        self.current_item = None

    def encodeData(self):
        mime_data = QByteArray()
        data_stream = QDataStream(mime_data, QIODevice.WriteOnly)

        selected_rows = self.selectedIndexes()
        for index in selected_rows:
            text = self.itemFromIndex(index).text()
            data_stream.writeInt(index.row())
            data_stream.writeInt(index.column())
            data_stream.writeQString(text)

        return mime_data

    def decodeData(self, data_stream):
        row = data_stream.readInt()
        column = data_stream.readInt()
        text = data_stream.readQString()

        return row, column, text


class especif_arquiv(QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setGeometry(10, 10, 500, 500)

        self.tipo_sec_label = QLabel("Tipo de Sessão", self)
        self.tipo_sec_label.setGeometry(0, 0, 100, 25)
        self.escolh_sec = QComboBox(self)
        self.escolh_sec.addItems(['Plenária', '1ª Câmara', '2ª Câmara'])
        self.escolh_sec.setGeometry(110, 0, 100, 25)

        self.escolh_data_label = QLabel("Data da Sessão", self)
        self.escolh_data_label.setGeometry(0, 30, 100, 25)
        self.escolh_data = QLineEdit(self)
        self.escolh_data.setPlaceholderText("Ex.: 25/04/2023")
        self.escolh_data.setGeometry(110, 30, 100, 25)

        self.presentes_label = QLabel("Presentes na Sessão", self)
        self.presentes_label.setGeometry(75, 60, 200, 25)
        self.presentes_label.setStyleSheet("font-weight: bold")

        self.anselmo_label = QLabel("Anselmo Roberto de Almeida Brito", self)
        self.anselmo_label.setGeometry(0, 90, 200, 25)
        self.anselmo_cond = QComboBox(self)
        self.anselmo_cond.addItems(['Escolha, se presente', 'Votante', 'Presidente', 'Presidente em exercício', 'Presente'])
        self.anselmo_cond.setGeometry(175, 90, 135, 25)

        self.fernando_label = QLabel("Fernando Ribeiro Toledo", self)
        self.fernando_label.setGeometry(0, 120, 200, 25)
        self.fernando_cond = QComboBox(self)
        self.fernando_cond.addItems(['Escolha, se presente', 'Votante', 'Presidente', 'Presidente em exercício', 'Presente'])
        self.fernando_cond.setGeometry(175, 120, 135, 25)

        self.maria_label = QLabel("Maria Cleide Costa Beserra", self)
        self.maria_label.setGeometry(0, 150, 200, 25)
        self.maria_cond = QComboBox(self)
        self.maria_cond.addItems(['Escolha, se presente', 'Votante', 'Presidente', 'Presidente em exercício', 'Presente'])
        self.maria_cond.setGeometry(175, 150, 135, 25)

        self.otavio_label = QLabel("Otávio Lessa de Geraldo Santos", self)
        self.otavio_label.setGeometry(0, 180, 200, 25)
        self.otavio_cond = QComboBox(self)
        self.otavio_cond.addItems(['Escolha, se presente', 'Votante', 'Presidente', 'Presidente em exercício', 'Presente'])
        self.otavio_cond.setGeometry(175, 180, 135, 25)

        self.renata_label = QLabel("Renata Pereira Pires Calheiros", self)
        self.renata_label.setGeometry(0, 210, 200, 25)
        self.renata_cond = QComboBox(self)
        self.renata_cond.addItems(['Escolha, se presente', 'Votante', 'Presidente', 'Presidente em exercício', 'Presente'])
        self.renata_cond.setGeometry(175, 210, 135, 25)

        self.rodrigo_label = QLabel("Rodrigo Siqueira Cavalcante", self)
        self.rodrigo_label.setGeometry(0, 240, 200, 25)
        self.rodrigo_cond = QComboBox(self)
        self.rodrigo_cond.addItems(['Escolha, se presente', 'Votante', 'Presidente', 'Presidente em exercício', 'Presente'])
        self.rodrigo_cond.setGeometry(175, 240, 135, 25)

        self.rosa_label = QLabel("Rosa Maria Ribeiro de Albuquerque", self)
        self.rosa_label.setGeometry(0, 270, 200, 25)
        self.rosa_cond = QComboBox(self)
        self.rosa_cond.addItems(['Escolha, se presente', 'Votante', 'Presidente', 'Presidente em exercício', 'Presente'])
        self.rosa_cond.setGeometry(175, 270, 135, 25)

        self.alberto_label = QLabel("Alberto Pires Alves de Abreu", self)
        self.alberto_label.setGeometry(0, 300, 200, 25)
        self.alberto_cond = QComboBox(self)
        self.alberto_cond.addItems(['Escolha, se presente', 'Votante', 'Presidente', 'Presidente em exercício', 'Presente'])
        self.alberto_cond.setGeometry(175, 300, 135, 25)

        self.sergio_label = QLabel("Sérgio Ricardo Maciel", self)
        self.sergio_label.setGeometry(0, 330, 200, 25)
        self.sergio_cond = QComboBox(self)
        self.sergio_cond.addItems(['Escolha, se presente', 'Votante', 'Presidente', 'Presidente em exercício', 'Presente'])
        self.sergio_cond.setGeometry(175, 330, 135, 25)

        self.mpc_label = QLabel("Membro do MPC", self)
        self.mpc_label.setGeometry(0, 360, 80, 25)
        self.mpc = QComboBox(self)
        self.mpc.addItems(['Stella de Barros Lima Mero Cavalcante', 'Ricardo Schneider Rodrigues', 'Pedro Barbosa Neto',
                           'Rafael Rodrigues de Alcântara', 'Enio Andrade Pimenta', 'Gustavo Henrique Albuquerque Santos'])
        self.mpc.setGeometry(85, 360, 225, 25)

        self.escolh_assinante_label = QLabel("Assinante", self)
        self.escolh_assinante_label.setGeometry(0, 390, 60, 25)
        self.assinante = QComboBox(self)
        self.assinante.addItems(['Lucas Nunes Aureliano Silva', 'Jéssica Luana Silva de Lima',
                                 'André Henrique da Rocha Alencar Rego'])
        self.assinante.setGeometry(85, 390, 225, 25)

        self.cargo_assin_label = QLabel("Cargo do Assinante", self)
        self.cargo_assin_label.setGeometry(0, 420, 100, 25)
        self.cargo_assin = QLineEdit(self)
        self.cargo_assin.setGeometry(110, 420, 100, 25)

        self.matric_ass_lab = QLabel("Matrícula do Assinante", self)
        self.matric_ass_lab.setGeometry(0, 450, 125, 25)
        self.matric_ass = QLineEdit(self)
        self.matric_ass.setGeometry(110, 450, 100, 25)


class AppDemo(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setFixedSize(1050, 525)
        self.setWindowTitle("Formatador de publicações. Closed Beta 1.0")
        self.lstbox_view = ListboxWidget(self)
        self.secao_view = especif_arquiv(self)

        self.btn_delete = QPushButton('Remover arquivo', self)
        self.btn_delete.setGeometry(400, 280, 200, 50)
        self.btn_delete.clicked.connect(self.delete_item)

        self.directory_btn = QPushButton('Escolha o destino', self)
        self.directory_btn.setGeometry(600, 280, 200, 50)
        self.directory_btn.clicked.connect(self.getDirectory)

        self.juntar_btn = QPushButton('Gerar arquivo para publicação', self)
        self.juntar_btn.setGeometry(800, 280, 200, 50)
        self.juntar_btn.clicked.connect(self.juntar_arq)

        self.arraste_aqui = QLabel("Arraste os arquivos na seção abaixo (apenas .docx ou .doc).", self)
        self.arraste_aqui.setGeometry(550, 65, 500, 50)

    def mes_do_ano(self, data):
        meses = {
            "01": "JANEIRO",
            "02": "FEVEREIRO",
            "03": "MARÇO",
            "04": "ABRIL",
            "05": "MAIO",
            "06": "JUNHO",
            "07": "JULHO",
            "08": "AGOSTO",
            "09": "SETEMBRO",
            "10": "OUTUBRO",
            "11": "NOVEMBRO",
            "12": "DEZEMBRO"
        }
        data = data.split("/")
        return f"{data[0]} DE {meses[data[1]]} DE {data[2]}"


    def juntar_arq(self, filename_master=None):
        logger = logging.getLogger(__name__)
        logger.setLevel(logging.DEBUG)

        formatter = logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s')
        file_handler = logging.FileHandler('debug.log')
        file_handler.setLevel(logging.DEBUG)
        file_handler.setFormatter(formatter)
        logger.addHandler(file_handler)

        try:
            logger.info('Começando a função juntar_arq')
            if not self.secao_view.escolh_data.text() or not self.secao_view.matric_ass or not self.secao_view.cargo_assin:
                QMessageBox.warning(self, "Atenção", "Por favor, preencha todos os campos antes de continuar.")
                logger.warning('Campos obrigatórios não preenchidos')
            else:
                items = [self.lstbox_view.item(i).text() for i in range(self.lstbox_view.count())]
                number_of_sections = len(items)
                if filename_master:
                    master = Document(filename_master)
                    logger.info('Arquivo mestre carregado com êxito')
                else:
                    master = Document()
                    logger.info('Criado novo arquivo mestre')
                style = master.styles['Normal']
                font = style.font
                font.name = 'Times New Roman'
                font.size = Pt(12)
                section = master.sections[0]
                section.left_margin = Cm(2)
                section.right_margin = Cm(2)
                section.top_margin = Cm(2)
                section.bottom_margin = Cm(2)
                composer = Composer(master)
                if self.secao_view.escolh_sec.currentText() == 'Plenária':
                    sessao = "PLENÁRIA"
                elif self.secao_view.escolh_sec.currentText() == '1ª Câmara':
                    sessao = "DA PRIMEIRA CÂMARA"
                elif self.secao_view.escolh_sec.currentText() == '2ª Câmara':
                    sessao = "DA SEGUNDA CÂMARA"
                logger.debug('Sessão escolhida: %s', sessao)

                dia = self.mes_do_ano(self.secao_view.escolh_data.text())
                logger.debug('Dia escolhido: %s', dia)

                cabecalho = master.add_paragraph()
                cabecalho.add_run('A CONSELHEIRA SUBSTITUTA DO TRIBUNAL DE CONTAS DE ALAGOAS, ').bold = False
                cabecalho.add_run(f"ANA RAQUEL RIBEIRO SAMPAIO CALHEIROS").bold = True
                cabecalho.add_run(f", NA SESSÃO ").bold = False
                cabecalho.add_run(f"{sessao}").bold = True
                cabecalho.add_run(f" DO DIA ").bold = False
                cabecalho.add_run(f"{dia}").bold = True
                cabecalho.add_run(', relatou os seguintes processos:').bold = False
                logger.debug('cabeçalho')

                sexo = {
                    "Rosa Maria Ribeiro de Albuquerque": "Conselheira",
                    "Renata Pereira Pires Calheiros": "Conselheira",
                    "Maria Cleide Costa Beserra": "Conselheira",
                    "Rodrigo Siqueira Cavalcante": "Conselheiro",
                    "Otávio Lessa de Geraldo Santos": "Conselheiro",
                    "Fernando Ribeiro Toledo": "Conselheiro",
                    "Anselmo Roberto de Almeida Brito": "Conselheiro",
                    "Alberto Pires Alves de Abreu": "Conselheiro Substituto",
                    "Sérgio Ricardo Maciel": "Conselheiro Substituto",
                    "Stella de Barros Lima Mero Cavalcante": "Procuradora",
                    "Ricardo Schneider Rodrigues": "Procurador",
                    "Pedro Barbosa Neto": "Procurador",
                    "Rafael Rodrigues de Alcântara": "Procurador",
                    "Enio Andrade Pimenta": "Procurador",
                    "Gustavo Henrique Albuquerque Santos": "Procurador"
                }

                if self.secao_view.anselmo_cond.currentText() == "Presidente":
                    presidente = "Anselmo Roberto de Almeida Brito"
                elif self.secao_view.fernando_cond.currentText() == "Presidente":
                    presidente = "Fernando Ribeiro Toledo"
                elif self.secao_view.maria_cond.currentText() == "Presidente":
                    presidente = "Maria Cleide Costa Beserra"
                elif self.secao_view.otavio_cond.currentText() == "Presidente":
                    presidente = "Otávio Lessa de Geraldo Santos"
                elif self.secao_view.renata_cond.currentText() == "Presidente":
                    presidente = "Renata Pereira Pires Calheiros"
                elif self.secao_view.rodrigo_cond.currentText() == "Presidente":
                    presidente = "Rodrigo Siqueira Cavalcante"
                elif self.secao_view.rosa_cond.currentText() == "Presidente":
                    presidente = "Rosa Maria Ribeiro de Albuquerque"
                elif self.secao_view.alberto_cond.currentText() == "Presidente":
                    presidente = "Alberto Pires Alves de Abreu"
                elif self.secao_view.sergio_cond.currentText() == "Presidente":
                    presidente = "Sérgio Ricardo Maciel"
                else:
                    presidente = False

                logger.debug('presidente: %s', presidente)

                if self.secao_view.anselmo_cond.currentText() == "Presidente em exercício":
                    presidente_em_exerc = "Anselmo Roberto de Almeida Brito"
                elif self.secao_view.fernando_cond.currentText() == "Presidente em exercício":
                    presidente_em_exerc = "Fernando Ribeiro Toledo"
                elif self.secao_view.maria_cond.currentText() == "Presidente em exercício":
                    presidente_em_exerc = "Maria Cleide Costa Beserra"
                elif self.secao_view.otavio_cond.currentText() == "Presidente em exercício":
                    presidente_em_exerc = "Otávio Lessa de Geraldo Santos"
                elif self.secao_view.renata_cond.currentText() == "Presidente em exercício":
                    presidente_em_exerc = "Renata Pereira Pires Calheiros"
                elif self.secao_view.rodrigo_cond.currentText() == "Presidente em exercício":
                    presidente_em_exerc = "Rodrigo Siqueira Cavalcante"
                elif self.secao_view.rosa_cond.currentText() == "Presidente em exercício":
                    presidente_em_exerc = "Rosa Maria Ribeiro de Albuquerque"
                elif self.secao_view.alberto_cond.currentText() == "Presidente em exercício":
                    presidente_em_exerc = "Alberto Pires Alves de Abreu"
                elif self.secao_view.sergio_cond.currentText() == "Presidente em exercício":
                    presidente_em_exerc = "Sérgio Ricardo Maciel"
                else:
                    presidente_em_exerc = False

                logger.debug('presidente em exerc: %s', presidente_em_exerc)

                votantes = []
                if self.secao_view.anselmo_cond.currentText() == "Votante":
                    votantes.append("Anselmo Roberto de Almeida Brito")
                if self.secao_view.fernando_cond.currentText() == "Votante":
                    votantes.append("Fernando Ribeiro Toledo")
                if self.secao_view.maria_cond.currentText() == "Votante":
                    votantes.append("Maria Cleide Costa Beserra")
                if self.secao_view.otavio_cond.currentText() == "Votante":
                    votantes.append("Otávio Lessa de Geraldo Santos")
                if self.secao_view.renata_cond.currentText() == "Votante":
                    votantes.append("Renata Pereira Pires Calheiros")
                if self.secao_view.rodrigo_cond.currentText() == "Votante":
                    votantes.append("Rodrigo Siqueira Cavalcante")
                if self.secao_view.rosa_cond.currentText() == "Votante":
                    votantes.append("Rosa Maria Ribeiro de Albuquerque")
                if self.secao_view.alberto_cond.currentText() == "Votante":
                    votantes.append("Alberto Pires Alves de Abreu")
                if self.secao_view.sergio_cond.currentText() == "Votante":
                    votantes.append("Sérgio Ricardo Maciel")

                logger.debug('votantes: %s', votantes)

                presente = []

                if self.secao_view.anselmo_cond.currentText() == "Presente":
                    presente.append("Anselmo Roberto de Almeida Brito")
                if self.secao_view.fernando_cond.currentText() == "Presente":
                    presente.append("Fernando Ribeiro Toledo")
                if self.secao_view.maria_cond.currentText() == "Presente":
                    presente.append("Maria Cleide Costa Beserra")
                if self.secao_view.otavio_cond.currentText() == "Presente":
                    presente.append("Otávio Lessa de Geraldo Santos")
                if self.secao_view.renata_cond.currentText() == "Presente":
                    presente.append("Renata Pereira Pires Calheiros")
                if self.secao_view.rodrigo_cond.currentText() == "Presente":
                    presente.append("Rodrigo Siqueira Cavalcante")
                if self.secao_view.rosa_cond.currentText() == "Presente":
                    presente.append("Rosa Maria Ribeiro de Albuquerque")
                if self.secao_view.alberto_cond.currentText() == "Presente":
                    presente.append("Alberto Pires Alves de Abreu")
                if self.secao_view.sergio_cond.currentText() == "Presente":
                    presente.append("Sérgio Ricardo Maciel")

                logger.debug('presentes: %s', presente)

                if presidente:
                    pres = presidente
                    pres1 = "Presidente"
                else:
                    pres = presidente_em_exerc
                    pres1 = "Presidente em exercício"

                logger.debug('pos presidente: %s', pres)
                logger.debug('nome presidente: %s', pres1)

                for i in range(number_of_sections):
                    doc_temp = Document(items[i])
                    for paragraph in doc_temp.paragraphs:
                        paragraph.style = doc_temp.styles['Normal']
                        paragraph.paragraph_format.space_before = Pt(0)
                        paragraph.paragraph_format.space_after = Pt(0)
                    composer.append(doc_temp)

                    logger.debug('composer append(doc tempo) executado')

                    presentes = master.add_paragraph()
                    presentes.add_run(f"\nConselheira Substituta ").bold = False
                    presentes.add_run(f"Ana Raquel Ribeiro Sampaio Calheiros ").bold = True
                    presentes.add_run(f"- Relatora").bold = False

                    logger.debug('paragrafo cons ana raquel executado')

                    presentes.add_run(f"\n{sexo[pres]} ").bold = False
                    presentes.add_run(f"{pres}").bold = True
                    presentes.add_run(f" - {pres1}").bold = False

                    logger.debug('paragrafo presidente executado')

                    presentes.add_run(f"\nTomaram parte na votação:").bold = True

                    logger.debug('paragrafo tomaram parte executado')

                    for votante in votantes:
                        presentes.add_run(f"\n{sexo[votante]} ").bold = False
                        presentes.add_run(f"{votante}").bold = True

                    logger.debug('paragrafo votantes executado')

                    for presen in presente:
                        presentes.add_run(f"\n{sexo[presen]} ").bold = False
                        presentes.add_run(f"{presen}").bold = True
                        presentes.add_run(f" - Presente").bold = False

                    logger.debug('paragrafo presente executado')

                    presentes.add_run(f"\n{sexo[self.secao_view.mpc.currentText()]} ").bold = False
                    presentes.add_run(f"{self.secao_view.mpc.currentText()}").bold = True
                    presentes.add_run(f" - Ministério Público de Contas - Presente").bold = False

                    logger.debug('paragrafo MPC executado')

                assinante = self.secao_view.assinante.currentText()
                cargo_assin = self.secao_view.cargo_assin.text()
                mat_assin = self.secao_view.matric_ass.text()

                assinatura = master.add_paragraph()
                assinatura.add_run(f"\n{assinante}").bold = True
                assinatura.add_run(f"\n{cargo_assin}\nMatrícula {mat_assin}").bold = False
                assinatura.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                logger.debug('paragrafo assinatura final executado')

                composer.save("combined_file.docx")
                QMessageBox.information(self, "Sucesso", "Arquivo gerado com sucesso!")

        except Exception as e:
            logger.exception("Ocorreu um erro: %s", e)

    def getDirectory(self):
        folder = QFileDialog.getExistingDirectory(
            self,
            caption='Selecione o destino'
        )
        print(folder)

    def delete_item(self):
        for item in self.lstbox_view.selectedItems():
            self.lstbox_view.takeItem(self.lstbox_view.row(item))


app = QApplication(sys.argv)
demo = AppDemo()
demo.show()
sys.exit(app.exec())